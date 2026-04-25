# -*- coding: utf-8 -*-
"""
LLM Dataset Builder
===================
Trích xuất thông tin (Metadata) từ các file `.docx` mẫu và chuyển đổi thành
tập dữ liệu (Dataset) chuẩn JSONL để huấn luyện mô hình ngôn ngữ lớn (LLM).
"""

import os
import sys
import json
import glob

# Fix python path for importing src from the root directory
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from docx import Document
from tqdm import tqdm

from src.config import Config

class DatasetBuilder:
    def __init__(self, data_dir=Config.RAW_DOCX_DIR, output_dir=Config.LLM_TRAINING_DIR):
        self.data_dir = data_dir
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        
    def extract_text_from_docx(self, docx_path):
        """Đọc toàn bộ text từ file Word."""
        try:
            doc = Document(docx_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(f"[TABLE] {cell.text.strip()}")
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error reading {docx_path}: {e}")
            return None

    def construct_synthetic_json_label(self, text):
        """
        Nội suy nhãn JSON từ văn bản (.docx) sạch.
        Sử dụng Regex nâng cao để bóc xuất chính xác các trường thông tin
        từ nhiều loại định dạng văn bản hành chính Việt Nam.
        """
        import re
        
        so_hieu = ""
        ngay_ban_hanh = ""
        trich_yeu = ""
        nguoi_ky = ""
        co_quan_ban_hanh = ""
        
        full_text = text  # Giữ nguyên toàn bộ text để regex tìm sâu
        lines = text.split('\n')
        
        # --- 1. Số hiệu: Tìm pattern "Số: XXX/YY-ZZ" ---
        so_match = re.search(r'Số[:\s]+(\d+\s*/\s*[\w\-\.]+)', full_text)
        if so_match:
            so_hieu = "Số: " + so_match.group(1).strip()
        
        # --- 2. Ngày ban hành (đa dạng format) ---
        # Format 1: "ngày DD/MM/YYYY" hoặc "ngày DD tháng MM năm YYYY"
        date_match = re.search(
            r'ngày\s+(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{4})', full_text, re.IGNORECASE
        )
        if date_match:
            ngay_ban_hanh = date_match.group(1).strip()
        else:
            # Format 2: "Ngày XX tháng YY năm ZZZZ"
            date_match2 = re.search(
                r'[Nn]gày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})',
                full_text
            )
            if date_match2:
                ngay_ban_hanh = f"{date_match2.group(1)}/{date_match2.group(2)}/{date_match2.group(3)}"
            else:
                # Format 3: Tìm ngày đầu tiên dạng DD/MM/YYYY bất kỳ trong 20 dòng đầu
                for line in lines[:20]:
                    dm = re.search(r'(\d{1,2}/\d{1,2}/\d{4})', line)
                    if dm:
                        ngay_ban_hanh = dm.group(1)
                        break
        
        # --- 3. Trích yếu: Tìm "V/v", "Về việc", hoặc dòng sau "QUYẾT ĐỊNH" ---
        for i, line in enumerate(lines[:20]):
            if "V/v" in line or "Về việc" in line:
                trich_yeu = line.strip()
                break
        if not trich_yeu:
            for i, line in enumerate(lines[:20]):
                if "QUYẾT ĐỊNH" in line.upper() or "THÔNG BÁO" in line.upper():
                    # Lấy dòng tiếp theo không rỗng làm trích yếu
                    for j in range(i+1, min(i+5, len(lines))):
                        if lines[j].strip() and len(lines[j].strip()) > 10:
                            trich_yeu = lines[j].strip()
                            break
                    break
        
        # --- 4. Cơ quan ban hành: Tìm tên IN HOA có "ỦY BAN", "SỞ", "BỘ", v.v. ---
        for line in lines[:10]:
            upper_line = line.strip().upper()
            if any(kw in upper_line for kw in ["ỦY BAN", "SỞ ", "BỘ ", "TRƯỜNG ", "UBND", "ĐẠI HỌC"]):
                co_quan_ban_hanh = line.strip()
                break
        
        # --- 5. Người ký: Tìm tên IN HOA ở phần cuối văn bản ---
        for line in lines[-15:]:
            clean = line.strip()
            if clean.isupper() and 5 < len(clean) < 40 and " " in clean:
                # Loại trừ các tiêu đề "CỘNG HÒA...", chỉ lấy tên người
                if not any(kw in clean for kw in ["CỘNG HÒA", "ĐỘC LẬP", "HẠNH PHÚC", "VIỆT NAM"]):
                    nguoi_ky = clean

        return json.dumps({
            "so_hieu": so_hieu,
            "ngay_ban_hanh": ngay_ban_hanh,
            "trich_yeu": trich_yeu,
            "co_quan_ban_hanh": co_quan_ban_hanh,
            "nguoi_ky": nguoi_ky
        }, ensure_ascii=False)

    def build_vqa_dataset(self):
        """Build file train_vqa.jsonl chuẩn ShareGPT cho VLM (Qwen2-VL / LLaVA)."""
        docx_files = glob.glob(os.path.join(self.data_dir, "**/*.docx"), recursive=True)
        if not docx_files:
            print(f"⚠️ Không tìm thấy file docx tại {self.data_dir}. Sinh data giả lập để test script.")
            # Mock data để script train không bị rỗng
            docx_files = ["mock1.docx", "mock2.docx"]
            
        output_file = os.path.join(self.output_dir, "train_vqa.jsonl")
        
        system_prompt = "Hãy trích xuất các thông tin hành chính từ văn bản trong ảnh và trả về định dạng JSON nghiêm ngặt."
        
        records_created = 0
        with open(output_file, 'w', encoding='utf-8') as f:
            for file_path in tqdm(docx_files, desc="Building VQA Dataset for Qwen2-VL"):
                if "mock" in file_path:
                    # Mock logic
                    text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n\nSố: 123/QĐ-UBND\nNgày 01 tháng 01 năm 2026\n\nQUYẾT ĐỊNH\nVề việc nâng cấp hệ thống AI\n\nCHỦ TỊCH\nNGUYỄN VĂN A"
                    image_path = "data/processed/stamped_images/mock1_page_1.png"
                else:
                    text = self.extract_text_from_docx(file_path)
                    # Mapping logic to the converted stamped images (Phase 1)
                    parent_name = os.path.basename(os.path.dirname(file_path))
                    stem = os.path.splitext(os.path.basename(file_path))[0]
                    # PDF -> Image usually extracts page 1
                    image_name = f"{parent_name}_{stem}_page_1.png"
                    image_path = f"data/processed/stamped_images/{image_name}"
                    
                if not text:
                    continue
                    
                json_label = self.construct_synthetic_json_label(text)
                
                # Định dạng ShareGPT/VQA chuẩn cho Qwen2-VL / LLaVA
                record = {
                    "id": f"vqa_doc_{records_created}",
                    "image": image_path,
                    "conversations": [
                        {
                            "from": "user",
                            "value": f"<image>\n{system_prompt}"
                        },
                        {
                            "from": "assistant",
                            "value": json_label
                        }
                    ]
                }
                f.write(json.dumps(record, ensure_ascii=False) + '\n')
                records_created += 1
                
        print(f"✅ Đã tạo thành công {records_created} mẫu huấn luyện VLM tại {output_file}")

if __name__ == "__main__":
    builder = DatasetBuilder()
    builder.build_vqa_dataset()
