import sys
from docx import Document

def read_doc(path):
    try:
        doc = Document(path)
        print(f"--- TEXT CONTENT ---")
        for p in doc.paragraphs:
            if p.text.strip():
                print(p.text.strip())
        print(f"\n--- TABLES ---")
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                print(" | ".join(row_data))
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == '__main__':
    read_doc(sys.argv[1])
